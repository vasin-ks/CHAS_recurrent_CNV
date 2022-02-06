#Name_off_pattient	2.3608007	GainMosaic	3	22268033	32989953	10721.927	9647	1111	p24.3	p22.3	5	2	UBE2E2-AS1, UBE2E2, UBE2E1-AS1, UBE2E1, NKIRAS1, UBE2E2 (602163), UBE2E1 (602916)


from docx import Document  #Install python-docx in PyCharm (I write to and run in it)
from docx.shared import RGBColor
document = Document()
blue = RGBColor(192, 192, 192)

from builtins import print
import re

inf = open("DataForComparison.bed").read() # DataForComparison.bed This file contains the coordinates of unreal patients. These are coordinates from the DGV base. They cannot be used. With them we compare the coordinates of the current investigated. The base should be supplemented by you independently !!!!
print("Enter the total number of CNVs from the CHAS program")
m = int(input())
print("Copy", m, "CNV from CHAS")
#p 15 columns from the program, ", "one line example:",
# "1:Name_off_pattient	2:2.3608007	3:GainMosaic	4:3	5:22268030	6:32989956	7:10721.927	8:9647	9:1111	10:p24.3	11:p22.3	12:5	13:2	14:UBE2E2-AS1, UBE2E2, UBE2E1-AS1, UBE2E1, NKIRAS1, 15:UBE2E2 (602163), UBE2E1 (602916)
# 1:PatientName 2:CN_State 3:Type 4:Chromosome 5:Min 6:Max 7:Size 8:M.Count 9:M.M.Dis 10:Cyto.Start 11:Cyto.End 12:Gene.Count 13:Omim.G.Count 14:Genes 15:Omim.Genes

CNV_without_genes = 0 #+
CNV_coincidence_in_two_coordinates = 0 #+
Unique_CNV = 0 #+
CNV_unique_in_one_coordinates = 0#+
CNV_with_genes = 0 #+
for i in range(m):
    i = input().split()
    b = len(i)
    if len(i) == 13: # cnv in which there are no genes, they have flax == 11, if there are genes, then 12 or 13
        CNV_without_genes += 1
    if len(i) > 13: # cnv in which the columns with genes are filled, that is, 12 and 13 lines
        CNV_with_genes  += 1
        SpisocVg = []# the list of all genes will be from now to now
        SpisocOMIMg = []# the list of all genes will be from now to now

        cit = []   # cit this is a chromosomal coordinate to write it correctly (according to the nomenclature)
        if i[9] == i[10]: # cit this is a chromosomal coordinate to write it correctly (according to the nomenclature)
            cit += [i[9]] # cit this is a chromosomal coordinate to write it correctly (according to the nomenclature)
        if i[9] != i[10]: # cit this is a chromosomal coordinate to write it correctly (according to the nomenclature)
            cit += i[9:11] # cit this is a chromosomal coordinate to write it correctly (according to the nomenclature)
        cit = ''.join(cit) #" ".join(str(x) for x in cit)  ген sec22b
        deldup = i[1]  #This is to indicate a deletion or duplication with one digit 1.0 by 1
        deldup = deldup[0] #This is to indicate a deletion or duplication with one digit 1.0 by 1
        corodkaja = (i[3],cit, "(", i[4], "_", i[5], ")", "x", deldup)
        corodkaja = ''.join(corodkaja)  # Final short entry
        if i[4] not in inf and i[5] not in inf:
            Unique_CNV += 1
            i = ' '.join(i)
            i = i.replace('(', '[OMIM:')
            i = i.replace(')', ']')
            p = document.add_paragraph(' ')
            p.add_run('Unique: ')
            p.add_run(i)
            p.add_run(' exon or intron ') #Using ucsc, you will later see which region of the gene is affected
            p.add_run(corodkaja)
        elif i[4] not in inf or i[5] not in inf:
            CNV_unique_in_one_coordinates += 1
            i = ' '.join(i)
            i = i.replace('(', '[OMIM:')
            i = i.replace(')', ']')
            p = document.add_paragraph(' ')
            p.add_run("Recurrent by 1 coordinates: ").font.color.rgb = RGBColor(255, 153, 51)
            p.add_run(i).font.color.rgb = RGBColor(255, 153, 51)
            p.add_run(' exon or intron   ').font.color.rgb = RGBColor(255, 153, 51)
            p.add_run(corodkaja).font.color.rgb = RGBColor(255, 153, 51)
        elif i[4] in inf and i[5] in inf:
            CNV_coincidence_in_two_coordinates += 1
            p = document.add_paragraph(" ")
            p.add_run("Recurrent: ").font.color.rgb = RGBColor(0, 102, 255)
            p.add_run("сoordinate ").font.color.rgb = RGBColor(0, 102, 255)
            p.add_run(str(inf.count(i[4]))).font.color.rgb = RGBColor(0, 102, 255)
            p.add_run(" meets ").font.color.rgb = RGBColor(0, 102, 255)
            p.add_run(i[4]).font.color.rgb = RGBColor(0, 102, 255)
            p.add_run(" time ").font.color.rgb = RGBColor(0, 102, 255)
            p.add_run(",and ").font.color.rgb = RGBColor(0, 102, 255)
            p.add_run(i[5]).font.color.rgb = RGBColor(0, 102, 255)
            p.add_run(" meets ").font.color.rgb = RGBColor(0, 102, 255)
            p.add_run(str(inf.count(i[5]))).font.color.rgb = RGBColor(0, 102, 255)
            p.add_run(" time ").font.color.rgb = RGBColor(0, 102, 255)
            p.add_run(" ".join(i)).font.color.rgb = RGBColor(0, 102, 255)



p = document.add_paragraph("Total CNV: ")
p.add_run(str(m))
p = document.add_paragraph("CNV without genes: ")
p.add_run(str(CNV_without_genes))
p = document.add_paragraph("Total CNV with genes: ")
p.add_run(str(CNV_with_genes))
p = document.add_paragraph("Recurrent CNV by 2 coordinates: ")
p.add_run(str(CNV_coincidence_in_two_coordinates))
p = document.add_paragraph("Total CNV +- unique: ")
p.add_run(str(Unique_CNV + CNV_unique_in_one_coordinates))
p = document.add_paragraph("Unique/NOT recurrent CNV: ")
p.add_run(str(Unique_CNV))
p = document.add_paragraph("CNV recurrent on 1 coordinate: ")
p.add_run(str(CNV_unique_in_one_coordinates))
#,i[2],i[8], i[9], "(", i[3], "_", i[4], ")", "x", "i[1]" sep="")

print(" CNV Catcher Finished")


document.save(r"C:\Users\Vasin\Desktop\CNV-catcher-result.doc")
#write your path to save the analysis results

